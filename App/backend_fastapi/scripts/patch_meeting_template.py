"""회의록 양식 docx에 작성 예시와 안내를 넣는다.

양식은 사내 서식이라 서식 자체를 유지해야 한다. 그래서 새로 만들지 않고 기존
파일의 셀 텍스트만 채운다. 실행한 뒤 반드시 실제 파일을 열어 눈으로 확인한다.

실행항목 칸(7행)의 1~3번 셀은 가로 병합되어 있어 python-docx가 같은 셀 객체를
그대로 반환한다. 그 안의 중첩 표도 물리적으로 하나뿐이므로 병합 셀 중 한 곳에서만
채운다 — 3번 모두 순회하면 예시 문구가 3번 겹쳐 써진다.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

TEMPLATE = (
    Path(__file__).resolve().parents[2]
    / "frontend" / "public" / "templates" / "meeting-minutes-template.docx"
)
EXAMPLE_ROW = "(예시) 박지수 · 로그인 오류 원인 파악 · 8/20"
ROW_HINT = "※ 항목이 더 필요하면 표에서 행을 추가하세요."
PARTICIPANT_HINT = "(쉼표로 구분: 박지수, 유소은)"

ACTION_ITEM_ROW = 7          # 실행항목 행 (Step 5에서 확인)
ACTION_ITEM_MERGED_CELL = 1  # 1~3번 셀은 가로 병합된 동일 셀. 하나만 쓴다
PARTICIPANT_ROW = 1          # 참석인원/참석자 행
PARTICIPANT_INPUT_CELL = 3   # 참석자 입력 칸


def _fill_existing_run(paragraph, text: str) -> None:
    """새 run을 add_run으로 붙이는 대신, 이미 있는 빈 입력용 run에 글자를 채운다.

    양식 자체가 그 방식으로 만들어져 있다(우선순위 체크 칸의 내용도 같은 패턴으로
    두 번째 run에 채워져 있다). add_run으로 새 run을 만들면 그 run은 기본 글꼴
    크기를 쓰게 되어 옆의 다른 입력 텍스트와 크기가 달라진다.
    """
    paragraph.runs[-1].text = text


def main() -> int:
    document = Document(TEMPLATE)
    table = document.tables[0]

    # 실행항목 예시 행 + 안내 문구
    action_cell = table.rows[ACTION_ITEM_ROW].cells[ACTION_ITEM_MERGED_CELL]
    nested = action_cell.tables[0]
    _fill_existing_run(nested.rows[1].cells[1].paragraphs[0], EXAMPLE_ROW)

    existing_hint_run = action_cell.paragraphs[-1].runs[0]  # 기존 "※ 해당 우선순위..." 안내
    hint_paragraph = action_cell.add_paragraph()
    hint_run = hint_paragraph.add_run(ROW_HINT)
    hint_run.font.size = existing_hint_run.font.size
    hint_run.font.italic = existing_hint_run.font.italic

    # 참석자 입력 칸 안내
    participant_cell = table.rows[PARTICIPANT_ROW].cells[PARTICIPANT_INPUT_CELL]
    _fill_existing_run(participant_cell.paragraphs[0], PARTICIPANT_HINT)

    document.save(TEMPLATE)
    print(f"패치 완료: {TEMPLATE}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
