"""회의록 양식의 실행항목 표를 4열(실행 항목 · 담당자 · 완료 기한 · 우선순위)로 바꾼다.

기존 2열 표는 우선순위를 `[ ] 긴급 [ ] 보통 [ ] 낮음` 체크로 고르고, 내용 칸 하나에
"누가 · 무엇을 · 언제까지"를 가운뎃점으로 이어 적게 했다. 체크가 번거롭고 가운뎃점을
빠뜨리면 담당자가 제목에 박힌다.

**우선순위 칸에 "보통"을 미리 인쇄하는 것이 이 표의 핵심 제약이다.** docx 추출기는
표를 셀마다 한 줄로 풀고 빈 셀은 줄 자체를 남기지 않아, 열 위치로는 칸을 구분할 수 없다.
우선순위 값(긴급/보통/낮음)이 행의 끝을 표시해야 담당자나 기한을 비운 행에서도 칸이
밀리지 않는다. 값을 비워두면 그 행은 다음 행과 합쳐진다.

양식은 사내 서식이라 서식 자체를 유지해야 하므로 새로 만들지 않고 기존 표를 고친다.
실행한 뒤 반드시 실제 파일을 열어 눈으로 확인한다.
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

TEMPLATE = (
    Path(__file__).resolve().parents[2]
    / "frontend" / "public" / "templates" / "meeting-minutes-template.docx"
)

ACTION_ITEM_MERGED_CELL = 1  # 1~3번 셀은 가로 병합된 동일 셀. 하나만 쓴다

# 상단 정보란에서 참석인원·참석자를 뺀다. 참석자는 업로드 화면에서 고르는 값이 단일
# 출처라 문서에 적어도 무시된다(MeetingTemplateParser 가 첫 헤딩 앞의 줄을 버리고,
# 파서 테스트가 그 동작을 못박아 뒀다). 두 곳에 적게 하면 어긋날 때 어느 쪽이 맞는지 모른다.
MEETING_TITLE_LABEL = "회의 제목"

COLUMN_HEADERS = ["실행 항목", "담당자", "완료 기한", "우선순위"]
# 기존 2열 표의 전체 폭(3240 x 2 = 6480 twip)을 그대로 나눈다. 표가 넓어지면 셀 밖으로 삐져나온다.
COLUMN_WIDTHS = [3000, 1100, 1200, 1180]

# 양식이 쓰던 글자 크기(w:sz 는 half-point 단위 - 16=8pt, 18=9pt).
# 열마다 출신이 달라 그냥 두면 실행 항목 칸만 한 단계 작게 나온다.
HEADER_FONT_SIZE = "16"
BODY_FONT_SIZE = "18"
DEFAULT_PRIORITY = "보통"
BLANK_ROWS = 5

# 예시는 표 안이 아니라 안내 줄에 둔다. 4열이라 예시가 네 칸으로 흩어지는데, 예시 행을
# 걸러내는 "(예시)" 접두사 규칙은 첫 칸만 잡는다. 나머지 세 칸이 살아남아 사용자가
# 예시를 지우지 않고 올리면 "박지수 / 8/20 / 긴급"이 유령 To-Do가 된다.
HINTS = [
    "※ 예: 로그인 오류 원인 파악 | 박지수 | 8/20 | 긴급",
    "※ 우선순위는 긴급 / 보통 / 낮음 중 하나입니다. 그대로 두면 보통으로 처리되니 지우지 마세요.",
    "※ 담당자나 기한이 정해지지 않았으면 비워두세요. 임의로 채우지 않습니다.",
    "※ 항목이 더 필요하면 표에서 행을 추가하세요.",
]


def _set_cell_text(cell, text: str, template_cell=None) -> None:
    """셀 글자를 바꾸되 글꼴은 표가 원래 쓰던 것을 따른다.

    양식의 라벨 칸은 "서식 없는 빈 run + 서식을 가진 run" 두 개로 되어 있다. 첫 run 에
    글자를 쓰면 글꼴 지정이 통째로 빠져 그 칸만 기본 글꼴로 보인다. 서식(rPr)을 가진
    run 을 찾아 거기에 쓰고, 나머지 run 은 지우지 않고 글자만 비운다 - run 을 지우면
    문단에 남은 서식 정보까지 함께 사라진다.
    """
    paragraph = cell.paragraphs[0]
    runs = list(paragraph.runs)
    if runs:
        target = next(
            (run for run in reversed(runs) if run._element.find(qn("w:rPr")) is not None),
            runs[-1],
        )
        for run in runs:
            if run is not target:
                run.text = ""
        target.text = text
        return

    run = paragraph.add_run(text)
    source = template_cell.paragraphs[0].runs if template_cell else []
    if source:
        run.font.size = source[0].font.size
        run.font.name = source[0].font.name


def _replace_nested_table(nested) -> None:
    """중첩 표를 4열로 다시 채운다. 행 서식은 기존 행을 복제해 유지한다."""
    template_row = copy.deepcopy(nested.rows[1]._tr)

    for row in list(nested.rows)[1:]:
        row._tr.getparent().remove(row._tr)

    _rewrite_grid(nested)

    header = nested.rows[0]
    _widen_row(header, len(COLUMN_HEADERS))
    for cell, text in zip(header.cells, COLUMN_HEADERS):
        _set_cell_text(cell, text)
    _normalize_font_size(header, HEADER_FONT_SIZE)

    for values in [["", "", "", DEFAULT_PRIORITY] for _ in range(BLANK_ROWS)]:
        nested._tbl.append(copy.deepcopy(template_row))
        row = nested.rows[-1]
        _widen_row(row, len(COLUMN_HEADERS))
        for cell, text in zip(row.cells, values):
            _set_cell_text(cell, text, template_cell=header.cells[0])
        _normalize_font_size(row, BODY_FONT_SIZE)


def _normalize_font_size(row, size: str) -> None:
    """한 행의 글자 크기를 맞춘다.

    4열은 옛 2열 표의 서로 다른 칸에서 복제돼 나온다. 체크 칸에서 온 열은 8pt,
    내용 칸에서 온 열은 9pt라 그냥 두면 사용자가 타이핑할 때 열마다 크기가 다르다.
    """
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                properties = run._element.find(qn("w:rPr"))
                if properties is None:
                    continue
                for tag in ("w:sz", "w:szCs"):
                    element = properties.find(qn(tag))
                    if element is not None:
                        element.set(qn("w:val"), size)


def _rewrite_grid(nested) -> None:
    """표의 열 정의(tblGrid)를 4열로 바꾼다.

    셀만 늘리고 이걸 안 고치면 Word 가 선언된 열 수만큼만 그려 표가 깨진다.
    """
    grid = nested._tbl.find(qn("w:tblGrid"))
    for column in list(grid.findall(qn("w:gridCol"))):
        grid.remove(column)
    for width in COLUMN_WIDTHS:
        column = grid.makeelement(qn("w:gridCol"), {qn("w:w"): str(width)})
        grid.append(column)


def _widen_row(row, columns: int) -> None:
    """2열 행을 4열로 늘리고 각 셀 폭을 열 정의에 맞춘다."""
    while len(row.cells) < columns:
        row._tr.append(copy.deepcopy(row.cells[-1]._tc))
        _set_cell_text(row.cells[-1], "")
    while len(row.cells) > columns:
        row.cells[-1]._tc.getparent().remove(row.cells[-1]._tc)
    for cell, width in zip(row.cells, COLUMN_WIDTHS):
        properties = cell._tc.get_or_add_tcPr()
        element = properties.find(qn("w:tcW"))
        if element is None:
            element = properties.makeelement(qn("w:tcW"), {})
            properties.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")


def _replace_hints(action_cell) -> None:
    """※ 안내 문구를 새 양식 기준으로 갈아 끼운다."""
    hint_paragraphs = [p for p in action_cell.paragraphs if p.text.strip().startswith("※")]
    sample = hint_paragraphs[0].runs[0] if hint_paragraphs and hint_paragraphs[0].runs else None

    for paragraph in hint_paragraphs:
        paragraph._element.getparent().remove(paragraph._element)

    for text in HINTS:
        paragraph = action_cell.add_paragraph()
        run = paragraph.add_run(text)
        run.font.italic = True
        run.font.size = sample.font.size if sample is not None else Pt(9)


def _row_index(table, label: str):
    """라벨 칸으로 행을 찾는다. 행을 지우면 번호가 밀리므로 번호를 박아두지 않는다."""
    for index, row in enumerate(table.rows):
        if row.cells[0].text.strip() == label:
            return index
    return None


def _trim_header_rows(table) -> None:
    """상단 정보란을 회의 제목 · 작성자 · 일시 · 장소 네 칸으로 줄인다.

    같은 스크립트를 두 번 돌려도 결과가 같도록, 이미 고친 라벨은 건너뛴다.
    """
    kind_row = _row_index(table, "회의종류")
    if kind_row is not None:
        _set_cell_text(table.rows[kind_row].cells[0], MEETING_TITLE_LABEL)

    attendee_row = _row_index(table, "참석인원")
    if attendee_row is not None:
        table.rows[attendee_row]._tr.getparent().remove(table.rows[attendee_row]._tr)


def main() -> int:
    document = Document(TEMPLATE)
    table = document.tables[0]

    _trim_header_rows(table)

    action_row = _row_index(table, "실행항목")
    action_cell = table.rows[action_row].cells[ACTION_ITEM_MERGED_CELL]

    _replace_nested_table(action_cell.tables[0])
    _replace_hints(action_cell)

    document.save(TEMPLATE)
    print(f"저장함: {TEMPLATE}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
