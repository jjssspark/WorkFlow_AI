"""회의록 양식(.docx)을 생성한다.

산출물: App/frontend/public/templates/meeting-minutes-template.docx

파일명을 ASCII로 두는 이유: 정적 서빙 경로에 한글이 들어가면 개발 서버와 nginx의
URL 인코딩 처리가 갈린다. 사용자가 보는 한글 파일명은 <a download> 속성이 지정한다.

양식은 사내 회의록 서식을 따른다 — 회색 라벨 열 + 격자 테두리 + 큰 본문 칸 + 하단 결재란.

라벨 칸(회의내용/결정사항/실행항목/비고)은 백엔드 파서(MeetingTemplateParser)가 인식하는
앵커다. 라벨 문구를 바꾸면 파서의 LABEL_HEADINGS도 함께 고쳐야 한다.

실행항목의 우선순위 체크박스도 백엔드가 읽는다. 표기를 바꾸면 FastAPI의
parse_action_items()도 함께 고쳐야 한다.

실행: python scripts/generate_meeting_template.py
의존성: pip install python-docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT_PATH = (
    Path(__file__).resolve().parents[1]
    / "App/frontend/public/templates/meeting-minutes-template.docx"
)

LABEL_SHADING = "D9D9D9"
PRIORITY_CHOICES = "[ ] 긴급    [ ] 보통    [ ] 낮음"

# 상단 정보란. 회의명·일시·참석자는 업로드 화면에서 따로 입력하므로 분석에는 쓰지 않는다.
# 문서를 사내 서식처럼 보이게 하는 용도이고, 파서는 이 구역을 통째로 무시한다.
INFO_ROWS = [
    ("회의종류", "", "작성자", ""),
    ("참석인원", "", "참석자", ""),
    ("일시", "", "장소", ""),
]

WIDE_INFO_ROWS = [
    ("회의주제", "회의를 진행하는 목적"),
    ("회의안건", "다룰 안건"),
]

BODY_ROWS = [
    ("회의내용", "회의 중 나온 논의를 자세히 적어주세요."),
    ("결정사항", "회의에서 확정된 내용을 적어주세요."),
]

REMARK_ROW = ("비고", "특이사항 · 리스크 · 다음 회의 일정을 적어주세요.")

ACTION_ITEM_ROWS = 4


def _shade(cell, color: str = LABEL_SHADING) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _write(cell, text: str, *, bold: bool = False, size: int = 9, italic: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def _label(cell, text: str) -> None:
    _shade(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _write(cell, text, bold=True)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _placeholder(cell, text: str) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _write(cell, text, italic=True, size=8)


def _add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("회 의 록")
    run.bold = True
    run.font.size = Pt(18)


def _add_main_grid(document: Document) -> None:
    total_rows = len(INFO_ROWS) + len(WIDE_INFO_ROWS) + len(BODY_ROWS) + 2  # +실행항목 +비고
    table = document.add_table(rows=total_rows, cols=4)
    table.style = "Table Grid"

    index = 0
    for left_label, left_value, right_label, right_value in INFO_ROWS:
        row = table.rows[index]
        _label(row.cells[0], left_label)
        _write(row.cells[1], left_value)
        _label(row.cells[2], right_label)
        _write(row.cells[3], right_value)
        index += 1

    for label, placeholder in WIDE_INFO_ROWS:
        row = table.rows[index]
        _label(row.cells[0], label)
        merged = row.cells[1].merge(row.cells[3])
        _placeholder(merged, placeholder)
        index += 1

    for label, placeholder in BODY_ROWS:
        row = table.rows[index]
        row.height = Cm(4.5)
        _label(row.cells[0], label)
        merged = row.cells[1].merge(row.cells[3])
        _placeholder(merged, placeholder)
        index += 1

    # 실행항목: 우선순위는 작성자가 체크하고, 내용은 문장으로 자유롭게 적는다.
    action_row = table.rows[index]
    _label(action_row.cells[0], "실행항목")
    action_cell = action_row.cells[1].merge(action_row.cells[3])
    _build_action_item_table(action_cell)
    index += 1

    remark_label, remark_placeholder = REMARK_ROW
    remark_row = table.rows[index]
    remark_row.height = Cm(3)
    _label(remark_row.cells[0], remark_label)
    _placeholder(remark_row.cells[1].merge(remark_row.cells[3]), remark_placeholder)


def _build_action_item_table(container_cell) -> None:
    container_cell.text = ""
    nested = container_cell.add_table(rows=1 + ACTION_ITEM_ROWS, cols=2)
    nested.style = "Table Grid"

    header = nested.rows[0]
    _write(header.cells[0], "우선순위", bold=True, size=8)
    _write(header.cells[1], "내용 (누가 · 무엇을 · 언제까지)", bold=True, size=8)
    _shade(header.cells[0])
    _shade(header.cells[1])

    for row in nested.rows[1:]:
        _write(row.cells[0], PRIORITY_CHOICES, size=8)
        _write(row.cells[1], "")

    hint = container_cell.add_paragraph()
    hint_run = hint.add_run("※ 해당 우선순위의 대괄호 안에 v 를 넣어주세요. 예: [v] 긴급")
    hint_run.italic = True
    hint_run.font.size = Pt(7)


def build_document() -> Document:
    document = Document()
    _add_title(document)
    _add_main_grid(document)
    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT_PATH)
    print(f"생성 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
